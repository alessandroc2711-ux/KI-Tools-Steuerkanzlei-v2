"""
KanzleiOptimierer – Effiziente Mandantenfallbearbeitung
LIVE VERSION MIT .ENV SUPPORT
Mit vollständiger Selbstbedienungs-Oberfläche
"""

import os
import json
from dotenv import load_dotenv
import fitz
from docx import Document
import streamlit as st
from pdf2image import convert_from_bytes
import pytesseract
from openai import OpenAI

# ======================================
# LOAD ENV
# ======================================
load_dotenv()
API_KEY = os.getenv("OPENAI_API_KEY")
if not API_KEY:
    raise ValueError("OPENAI_API_KEY nicht gefunden. Bitte .env prüfen.")
client = OpenAI(api_key=API_KEY)

# ======================================
# CONFIG
# ======================================
DICT_FILE = "dictionary.json"

DEFAULT_DICT = {
    "Alexander Codazzi": "MANDANT_A",
    "Siemens Energy AG": "ARBEITGEBER_1",
    "BMW AG": "FIRMA_1",
    "Allianz SE": "FIRMA_2",
    "Deutsche Bank AG": "BANK_1",
    "Upwork": "PLATTFORM_1",
    "Thomas Berger": "PERSON_1",
    "Julia Berger": "PERSON_2",
    "Blumenstraße 12, 80331 München": "ADRESSE_1",
    "Maxvorstadt 45, 80799 München": "ADRESSE_2",
    "Sonnenstraße 88, 80331 München": "ADRESSE_3",
    "BMW Leasing": "FAHRZEUGLEASING_1"
}

# ======================================
# STREAMLIT CONFIG
# ======================================
st.set_page_config(page_title="KanzleiOptimierer", layout="wide")

st.title("KanzleiOptimierer – Effiziente Mandantenfallbearbeitung")

st.markdown("""
Willkommen bei **KanzleiOptimierer**.

Dieses Tool ermöglicht eine **vollständig selbstbedienbare Nutzung**:

- Upload und Anonymisierung von PDF- und Word-Dokumenten
- Bearbeitbares, manuelles Anonymisierungs-Wörterbuch
- Automatisierte steuerrechtliche Analyse
""")

# ======================================
# DICTIONARY FUNCTIONS
# ======================================
def load_dictionary():
    if os.path.exists(DICT_FILE):
        with open(DICT_FILE, "r", encoding="utf-8") as f:
            dictionary = json.load(f)
    else:
        dictionary = DEFAULT_DICT.copy()

    return dict(sorted(dictionary.items()))


def save_dictionary(dictionary):
    sorted_dict = dict(sorted(dictionary.items()))

    with open(DICT_FILE, "w", encoding="utf-8") as f:
        json.dump(
            sorted_dict,
            f,
            indent=4,
            ensure_ascii=False
        )


def anonymize_text(text, dictionary):
    for original, replacement in dictionary.items():
        text = text.replace(original, replacement)

    return text


# ======================================
# FILE EXTRACTION
# ======================================
def extract_pdf_text(file_pdf):
    file_pdf.seek(0)
    content = file_pdf.read()

    pdf = fitz.open(
        stream=content,
        filetype="pdf"
    )

    text = "".join(
        page.get_text() for page in pdf
    )

    if text.strip():
        return text

    images = convert_from_bytes(content)

    text = "\n".join(
        pytesseract.image_to_string(
            img,
            lang="deu"
        )
        for img in images
    )

    return text


def extract_word_text(file_docx):
    file_docx.seek(0)

    doc = Document(file_docx)

    return "\n".join(
        para.text for para in doc.paragraphs
    )


# ======================================
# WORD EXPORT – OPTIMIERT
# ======================================
def create_word(text, filename="document_anon.docx"):
    doc = Document()

    lines = text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        # Heading
        if line.startswith("# "):
            doc.add_heading(
                line.replace("# ", ""),
                level=1
            )

        # Bullet list
        elif line.startswith("- "):
            doc.add_paragraph(
                line[2:],
                style="List Bullet"
            )

        # Numbered list
        elif len(line) > 2 and line[0].isdigit() and line[1] == ".":
            doc.add_paragraph(
                line,
                style="List Number"
            )

        # Markdown table
        elif line.startswith("|") and line.endswith("|"):
            table_lines = []

            while i < len(lines):
                current = lines[i].strip()

                if current.startswith("|") and current.endswith("|"):
                    table_lines.append(current)
                    i += 1
                else:
                    break

            cleaned_lines = [
                row for row in table_lines
                if "---" not in row
            ]

            if cleaned_lines:
                headers = [
                    cell.strip()
                    for cell in cleaned_lines[0].split("|")[1:-1]
                ]

                table = doc.add_table(
                    rows=1,
                    cols=len(headers)
                )

                hdr_cells = table.rows[0].cells

                for idx, header in enumerate(headers):
                    hdr_cells[idx].text = header

                for row_line in cleaned_lines[1:]:
                    row_cells = table.add_row().cells

                    values = [
                        cell.strip()
                        for cell in row_line.split("|")[1:-1]
                    ]

                    for idx, value in enumerate(values):
                        row_cells[idx].text = value

            continue

        # Empty line
        elif line == "":
            doc.add_paragraph("")

        # Normal text
        else:
            doc.add_paragraph(line)

        i += 1

    doc.save(filename)

    return filename


# ======================================
# PROMPT ENGINEERING
# ======================================
def build_tax_prompt(case_text):
    prompt = f"""
Du bist ein Senior-Steuerberater einer führenden deutschen Steuerkanzlei.

Erstelle eine steuerrechtliche Fachanalyse IMMER in exakt dem folgenden Format.
Die Struktur, Überschriften und Reihenfolge dürfen niemals verändert werden.

Verwende ausschließlich diese Formatvorlage:

# Executive Summary
Kurzfassung in 3–5 Sätzen.

# Steuerrechtliche Fragestellung
- Punkt 1
- Punkt 2
- Punkt 3

# Anwendbare Normen
| Norm | Relevanz | Bewertung |
|---|---|---|
| Paragraph | kurze Erklärung | hoch / mittel / niedrig |

# Relevante BFH / FG Urteile
| Urteil | Kernaussage | Relevanz für Fall |
|---|---|---|

# Expertenbewertung
Ausführliche fachliche Einschätzung in Fließtext.

# Risikoanalyse
| Bereich | Risiko | Begründung |
|---|---|---|
| Steuerlich | niedrig / mittel / hoch | Erklärung |
| Dokumentation | niedrig / mittel / hoch | Erklärung |
| Betriebsprüfung | niedrig / mittel / hoch | Erklärung |

# Handlungsempfehlung
1. Empfehlung
2. Empfehlung
3. Empfehlung

# Nächste Schritte für die Kanzlei
1. Sofortmaßnahme
2. Mandanten-Rückfrage
3. Dokumentationsschritt
4. Frist / Deadline

WICHTIGE REGELN:
- Immer exakt dieselben Überschriften verwenden
- Immer dieselbe Reihenfolge
- Tabellenformat IMMER beibehalten
- Keine zusätzlichen Kapitel
- Keine Abweichungen
- Antwort ausschließlich auf Deutsch
- professioneller Kanzlei-Stil
- juristisch präzise
- mandantenorientiert

Mandantenfall:
{case_text}
"""
    return prompt


# ======================================
# OPENAI CALL
# ======================================
def run_ai_analysis(prompt):
    response = client.responses.create(
        model="gpt-4.1",
        input=prompt
    )

    return response.output_text


# ======================================
# SESSION STATE
# ======================================
if "dictionary" not in st.session_state:
    st.session_state.dictionary = load_dictionary()

if "original_text" not in st.session_state:
    st.session_state.original_text = ""

if "anonymized_text" not in st.session_state:
    st.session_state.anonymized_text = ""


# ======================================
# SIDEBAR
# ======================================
st.sidebar.header("Anonymisierungs-Wörterbuch")

st.sidebar.markdown(
    "Format: `Original = Ersatz`\n\nBeispiel: `BMW AG = FirmaX`"
)

dictionary_text = "\n".join(
    f"{k} = {v}"
    for k, v in sorted(
        st.session_state.dictionary.items()
    )
)

edited_dictionary_text = st.sidebar.text_area(
    "Wörterbuch bearbeiten",
    value=dictionary_text,
    height=300
)

if st.sidebar.button("Wörterbuch speichern"):
    new_dict = {}

    for line in edited_dictionary_text.split("\n"):
        if "=" in line:
            key, value = line.split("=", 1)
            new_dict[key.strip()] = value.strip()

    st.session_state.dictionary = dict(
        sorted(new_dict.items())
    )

    save_dictionary(
        st.session_state.dictionary
    )

    if st.session_state.original_text:
        st.session_state.anonymized_text = anonymize_text(
            st.session_state.original_text,
            st.session_state.dictionary
        )

    st.sidebar.success(
        "Wörterbuch gespeichert und Text aktualisiert"
    )


# ======================================
# STEP 1 – FILE UPLOAD
# ======================================
st.header("1. Dokument Upload")

uploaded_file = st.file_uploader(
    "PDF oder Word auswählen",
    type=["pdf", "docx"]
)

if uploaded_file is not None and st.button("Dokument verarbeiten"):
    if uploaded_file.name.endswith(".pdf"):
        text = extract_pdf_text(uploaded_file)
    else:
        text = extract_word_text(uploaded_file)

    st.session_state.original_text = text

    st.session_state.anonymized_text = anonymize_text(
        text,
        st.session_state.dictionary
    )


# ======================================
# STEP 2 – TEXT
# ======================================
if st.session_state.original_text:
    st.subheader("Originaltext")

    st.text_area(
        "Original",
        st.session_state.original_text,
        height=200
    )

if st.session_state.anonymized_text:
    st.subheader("Anonymisierter Text")

    st.session_state.anonymized_text = st.text_area(
        "Bearbeitbarer Text – eine automatische Bearbeitung erfolgt basierend auf dem Anonymisierungs-Wörterbuch",
        st.session_state.anonymized_text,
        height=300
    )

    if st.button("Text erneut anonymisieren (Optional)"):
        st.session_state.anonymized_text = anonymize_text(
            st.session_state.original_text,
            st.session_state.dictionary
        )

        st.success("Text aktualisiert")

    filename = create_word(
        st.session_state.anonymized_text
    )

    with open(filename, "rb") as f:
        st.download_button(
            "Word herunterladen (Optional)",
            data=f,
            file_name=filename
        )


# ======================================
# STEP 3 – ANALYSE
# ======================================
if st.session_state.anonymized_text:
    st.header("2. Steuerliche Analyse")

    st.info(
        "Das Tool erstellt eine standardisierte steuerliche Auswertung des anonymisierten Falls."
    )

    if st.button("Analyse starten"):
        with st.spinner("Analyse läuft..."):
            prompt = build_tax_prompt(
                st.session_state.anonymized_text
            )

            result = run_ai_analysis(prompt)

        st.success("Analyse abgeschlossen")

        st.subheader("Analyseergebnis")
        st.markdown(result)

        filename_analysis = "Analyseergebnis.docx"

        create_word(
            result,
            filename=filename_analysis
        )

        with open(filename_analysis, "rb") as f:
            st.download_button(
                "Analyse als Word herunterladen",
                data=f,
                file_name=filename_analysis
            )
